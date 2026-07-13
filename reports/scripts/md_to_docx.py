"""
Convert a KuantorFlow report written in Markdown into a styled .docx.

Usage:
    python md_to_docx.py <report.md> [output.docx]

Requires: python-docx  (pip install python-docx)

Supported Markdown (the subset our reports use):
    # Title            ## Section (Heading 1)      ### Subsection (Heading 2)
    ---                horizontal rule (yellow underline)
    - bullets          1. numbered items           | tables | with | header |
    **bold**  *italic*  `code`  [link text](url)   (link text kept, shown in blue)
Hard-wrapped lines are joined; list continuation lines (indented) are appended
to their item.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x2E, 0x6B, 0xA0)
DARK = RGBColor(0x22, 0x30, 0x3C)
MUTED = RGBColor(0x5B, 0x6B, 0x7A)
YELLOW = "D9B13C"
BLUE_HEX = "2E6BA0"

# ---------------------------------------------------------------- parsing ---

INLINE = re.compile(
    r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`|\[[^\]\n]+?\]\([^)\n]+?\))"
)
LINK = re.compile(r"\[([^\]]+?)\]\(([^)]+?)\)")


def parse_inline(text):
    """Split text into runs: list of dicts {text, bold, italic, code, link}."""
    runs = []
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            runs.append({"text": part[2:-2], "bold": True})
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            runs.append({"text": part[1:-1], "italic": True})
        elif part.startswith("`") and part.endswith("`"):
            runs.append({"text": part[1:-1], "code": True})
        else:
            m = LINK.fullmatch(part)
            if m:
                runs.append({"text": m.group(1), "link": m.group(2)})
            else:
                runs.append({"text": part})
    return runs


def parse_markdown(text):
    """Parse report markdown into a list of blocks:
    ('title'|'h1'|'h2'|'p'|'bullet'|'num', runs) | ('table', rows) | ('hr',)
    where table rows are lists of run-lists (first row = header)."""
    blocks = []
    lines = text.splitlines()
    i, para = 0, []

    def flush_para():
        if para:
            blocks.append(("p", parse_inline(" ".join(para))))
            para.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_para()
        elif stripped == "---":
            flush_para()
            blocks.append(("hr",))
        elif stripped.startswith("### "):
            flush_para()
            blocks.append(("h2", parse_inline(stripped[4:])))
        elif stripped.startswith("## "):
            flush_para()
            blocks.append(("h1", parse_inline(stripped[3:])))
        elif stripped.startswith("# "):
            flush_para()
            blocks.append(("title", parse_inline(stripped[2:])))
        elif stripped.startswith("|"):
            flush_para()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    rows.append([parse_inline(c) for c in cells])
                i += 1
            blocks.append(("table", rows))
            continue
        elif stripped.startswith("- "):
            flush_para()
            item = stripped[2:]
            while i + 1 < len(lines) and lines[i + 1].startswith("  ") \
                    and lines[i + 1].strip() and not lines[i + 1].lstrip().startswith(("- ", "|")):
                i += 1
                item += " " + lines[i].strip()
            blocks.append(("bullet", parse_inline(item)))
        elif re.match(r"^\d+\.\s+", stripped):
            flush_para()
            item = re.sub(r"^\d+\.\s+", "", stripped)
            while i + 1 < len(lines) and lines[i + 1].startswith("  ") \
                    and lines[i + 1].strip() and not re.match(r"^\s*\d+\.\s+", lines[i + 1]):
                i += 1
                item += " " + lines[i].strip()
            blocks.append(("num", parse_inline(item)))
        else:
            para.append(stripped)
        i += 1
    flush_para()
    return blocks


# --------------------------------------------------------------- emitting ---

def add_runs(par, runs, base_size=11, base_color=DARK):
    for r in runs:
        run = par.add_run(r["text"])
        run.bold = r.get("bold", False)
        run.italic = r.get("italic", False)
        run.font.size = Pt(base_size)
        run.font.color.rgb = BLUE if r.get("link") else base_color
        if r.get("code"):
            run.font.name = "Consolas"


def add_rule(doc):
    par = doc.add_paragraph()
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), YELLOW)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    # proportional column widths from content length, 6.8" total
    weights = [max((len("".join(x["text"] for x in r[j])) if j < len(r) else 1)
                   for r in rows) or 1 for j in range(cols)]
    total = sum(weights)
    widths = [max(0.55, 6.8 * w / total) for w in weights]

    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = t.cell(i, j)
            cell.width = Inches(widths[j])
            runs = row[j] if j < len(row) else []
            par = cell.paragraphs[0]
            if i == 0:
                shade(cell, BLUE_HEX)
                add_runs(par, runs, base_size=10, base_color=RGBColor(0xFF, 0xFF, 0xFF))
                for r in par.runs:
                    r.bold = True
            else:
                add_runs(par, runs, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_docx(blocks, out_path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    for name, size, color in (("Heading 1", 15, BLUE), ("Heading 2", 12.5, DARK)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.83)
    sec.top_margin = sec.bottom_margin = Inches(0.75)

    for block in blocks:
        kind = block[0]
        if kind == "title":
            par = doc.add_paragraph()
            add_runs(par, block[1], base_size=24, base_color=BLUE)
            for r in par.runs:
                r.bold = True
            par.paragraph_format.space_after = Pt(4)
        elif kind == "h1":
            add_runs(doc.add_heading(level=1), block[1], base_size=15, base_color=BLUE)
        elif kind == "h2":
            add_runs(doc.add_heading(level=2), block[1], base_size=12.5)
        elif kind == "hr":
            add_rule(doc)
        elif kind == "p":
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(6)
            add_runs(par, block[1])
        elif kind == "bullet":
            par = doc.add_paragraph(style="List Bullet")
            par.paragraph_format.space_after = Pt(4)
            add_runs(par, block[1])
        elif kind == "num":
            par = doc.add_paragraph(style="List Number")
            par.paragraph_format.space_after = Pt(5)
            add_runs(par, block[1])
        elif kind == "table":
            add_table(doc, block[1])
    doc.save(out_path)


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    src = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    blocks = parse_markdown(src.read_text(encoding="utf-8"))
    build_docx(blocks, str(out))
    print(f"written: {out}")


if __name__ == "__main__":
    main()
